from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def create_requirements_doc(filename="TruthLens_Requirements.docx"):
    doc = Document()
    
    # Title
    title = doc.add_heading('TruthLens: Fake News Detection Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('System Requirements Specification (SRS)')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 1: Introduction
    doc.add_heading('1. Introduction', level=1)
    doc.add_heading('1.1 Purpose', level=2)
    doc.add_paragraph(
        "The purpose of this document is to detail the software and system requirements for "
        "TruthLens, an automated fake news detection platform. TruthLens is designed to ingest, "
        "process, and analyze news articles using Natural Language Processing (NLP) and Machine "
        "Learning (ML) to provide transparent, explainable predictions on article authenticity."
    )
    
    doc.add_heading('1.2 Scope', level=2)
    doc.add_paragraph(
        "The system will provide an end-to-end pipeline allowing users to:\n"
        "1. Upload and ingest text datasets (CSV, Excel, JSON).\n"
        "2. Clean and preprocess textual data using NLP (Stopwords, Lemmatization, Stemming).\n"
        "3. Visualize dataset metrics through an interactive dashboard.\n"
        "4. Train custom classification algorithms (e.g., Logistic Regression, Random Forest).\n"
        "5. Predict and verify the authenticity of new articles with Explainable AI (XAI).\n"
        "6. Generate system performance audits and reports."
    )
    
    # Section 2: Functional Requirements
    doc.add_heading('2. Functional Requirements', level=1)
    
    reqs = [
        ("FR-01: Data Ingestion", "The system shall allow users to upload datasets up to 200MB in CSV, XLSX, XLS, and JSON formats, and parse text features."),
        ("FR-02: NLP Preprocessing", "The system shall provide configurable NLP toggles to remove stopwords, apply WordNet Lemmatization, and apply Porter Stemming."),
        ("FR-03: Dashboard Analytics", "The system shall display exploratory data analysis (EDA) charts including class balance, word frequencies, and article word count distribution."),
        ("FR-04: Model Training", "The system shall allow users to select vectorization parameters (TF-IDF, N-grams) and train ML classifiers on the processed data."),
        ("FR-05: Real-time Detection", "The system shall evaluate raw text input against active models and return a confidence score alongside a 'Real' or 'Fake' classification."),
        ("FR-06: Explainable AI (XAI)", "The system shall highlight influential tokens in the input text to explain the classification decision."),
        ("FR-07: Audit Logging", "The system shall save training and prediction logs and allow users to export them as HTML reports.")
    ]
    for r_id, r_desc in reqs:
        p = doc.add_paragraph()
        p.add_run(r_id + ": ").bold = True
        p.add_run(r_desc)

    # Section 3: Non-Functional Requirements
    doc.add_heading('3. Non-Functional Requirements', level=1)
    
    n_reqs = [
        ("NFR-01: Usability", "The application shall feature a 'Flat Design' UI using a high-contrast light theme and intuitive workflow navigation."),
        ("NFR-02: Privacy & Security", "All data processing, model training, and inference must happen locally on the machine to guarantee data privacy. No external API calls are made for core ML operations."),
        ("NFR-03: Performance", "The system must process dataset uploads and complete preprocessing dry runs in under 10 seconds for sample batches."),
        ("NFR-04: Maintainability", "The system must be built using modular Python scripts and Streamlit multi-page routing to ensure future maintainability.")
    ]
    for r_id, r_desc in n_reqs:
        p = doc.add_paragraph()
        p.add_run(r_id + ": ").bold = True
        p.add_run(r_desc)

    # Section 4: Hardware & Software Requirements
    doc.add_heading('4. System Requirements', level=1)
    
    doc.add_heading('4.1 Software Requirements', level=2)
    doc.add_paragraph("• Language: Python 3.9+\n"
                      "• Frontend: Streamlit\n"
                      "• Machine Learning: Scikit-learn, Pandas, NumPy\n"
                      "• Natural Language Processing: NLTK\n"
                      "• Visualizations: Plotly\n"
                      "• Database: SQLite 3")

    doc.add_heading('4.2 Hardware Requirements', level=2)
    doc.add_paragraph("• Processor: Dual-core 2.0 GHz or higher\n"
                      "• RAM: 4 GB minimum (8 GB recommended for large datasets)\n"
                      "• Storage: 500 MB for application files and models (SSD recommended)\n"
                      "• Display: 1280x720 minimum resolution")
    
    doc.save(filename)
    print(f"Requirements document saved successfully as {filename}")

if __name__ == '__main__':
    create_requirements_doc()
